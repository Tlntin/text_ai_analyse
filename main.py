import PySimpleGUI as sg
from aip import AipNlp
from docx import Document
from docx.shared import RGBColor
import os
import pickle
import random
import time


class TextAIAnalyse(object):
    """
    :param
    """

    def __init__(self, doc_path, app_id, api_key, secret_key):
        """
        :param doc_path：文章路径
        :param app_id: 应用id，自己去百度ai控制台构建一个应用，就会有id了
        :param api_key:
        :param secret_key:
        """
        self.client = AipNlp(app_id, api_key, secret_key)
        self.document = Document(doc_path)
        self.doc_path = doc_path
        text_list1 = self.filter_style()
        self.text_list2 = self.filter_short_text(text_list1, 12)

    def filter_style(self):
        """
        样式过滤
        :param
        """
        delete_style = ['Title', 'Heading 1', 'Quote']  # 去除标题，一级标题，图表链接
        list1 = [x.text for x in self.document.paragraphs if x.style.name not in delete_style]
        return list1

    @staticmethod
    def filter_short_text(list1: list, length: int):
        """
        去除短文本
        :param list1: 列表格式的文本集
        :param length:最短文本长度
        """
        list2 = [x.strip() for x in list1]  # 去除两边空格
        list3 = [x for x in list2 if len(x) > length]
        return list3

    def split_text(self, list1: list):
        """
        对段落进行分句，粗糙分词
        :param
        """
        list2 = []
        for x in list1:
            x_list1 = x.split('。')  # 以句号进行分词,分号暂时不考虑
            for xx in x_list1:
                if xx[-1:] not in ['。', '；', '：']:  # 如果本局不是以句号结尾，则给它加上句号
                    xx += '。'
                    list2.append(xx)
        list3 = self.filter_short_text(list2, 10)
        return list3

    def split_text2(self, list1: list):
        """
        对段落进行分句，加上分号
        :param
        """
        list2 = []
        for x in list1:
            x_list1 = x.split('。')  # 以句号进行分词
            for xx in x_list1:
                x_list2 = xx.split('；')  # 以中文分号进行分词
                for xxx in x_list2:
                    if xxx[-1:] not in ['。', '；', '：']:  # 如果本局不是以句号结尾，则给它加上句号
                        xxx += '。'
                        list2.append(xxx)
        list3 = self.filter_short_text(list2, 10)
        return list3

    def ai_analyse(self, text1):
        """
        AI对句子进行纠错
        :param
        """
        result1 = None
        try:
            result1 = self.client.ecnet(text1)
        except Exception as err:
            return False
        vet = result1['item']['vec_fragment']  # 可替换词
        score = result1['item']['score']  # 评分
        if len(vet) == 0:
            return False  # 没有错误
        # elif score > 0.5:  # 如果可信度
        else:
            return result1  # 返回分析

    def save_analyse(self, result):
        """
        :param
        """
        text = result['text']
        text_encode = text.encode('gbk')
        vet_list = result['item']['vec_fragment']
        """ 开始写入word """
        basename = os.path.basename(self.doc_path)
        path = './分析结果/{}.docx'.format(basename[:6])
        dir_name = './分析结果'
        if not os.path.exists(dir_name):  # 如果文件夹不存在
            os.mkdir(dir_name)
        a = vet_list[0]['begin_pos']  # 获取第一个错误标签的开始位置
        b = vet_list[-1]['end_pos']  # 获取最后一个错误的结束位置
        if not os.path.exists(path):  # 如果文件不存在
            doc = Document()
        else:
            doc = Document(path)
        doc.add_paragraph('错误写法', style='heading 1')  # 一级标题
        p = doc.add_paragraph()  # 创建一个空段落
        p.add_run(text_encode[:a].decode('gbk'))  # 写入没有错误的部分
        start_list = [x['begin_pos'] for x in vet_list]  # 记录开始的位置
        end_list = [x['end_pos'] for x in vet_list]  # 记录结束的位置
        if len(vet_list) == 1:  # 如果只有一个错误
            run1 = p.add_run(text_encode[start_list[0]:end_list[0]].decode('gbk'))
            run1.font.color.rgb = RGBColor(255, 0, 0)  # 设置红色字体
        else:  # 如果有多个错误
            for i in range(len(vet_list)):
                run = p.add_run(text_encode[start_list[i]:end_list[i]].decode('gbk'))
                run.font.color.rgb = RGBColor(255, 0, 0)  # 设置红色字体
                if i < len(vet_list) - 1 and start_list[i + 1] - end_list[i] > 1:  # 如果后面一处错误与前面一处错误存在间距
                    p.add_run(text_encode[end_list[i]:start_list[i + 1]].decode('gbk'))  # 增加一个没有样式的普通字体
        if len(text_encode) - b > 1:  # 如果最后一个错误后面存在正确的内容
            p.add_run(text_encode[b:].decode('gbk'))  # 写入后面无错的内容
        doc.add_paragraph('正确写法', style='heading 1')  # 一级标题
        correct = result['item']['correct_query']  # 正确的内容
        doc.add_paragraph(correct)  # 写入正确的内容
        doc.save(path)


if __name__ == '__main__':
    my_font = 'Deja_Vu_Sans_Mono.ttf'
    my_font_style1 = (my_font, 11, "normal")
    # 菜单栏
    menu_def = [
        ['&菜单', ['使用说明', '更新记录']],
        ['&文件', ['载入配置', '保存配置']]
    ]
    # 布局栏
    layout1 = [
        [sg.Menu(menu_def, tearoff=True)],
        [sg.Text('APP_ID:', size=(12, None)), sg.Input(key='app_id')],
        [sg.Text('API_KEY', size=(12, None)), sg.Input(key='api_key')],
        [sg.Text('SECRET_KEY', size=(12, None)), sg.Input(key='secret_key')],
        [sg.Text('文件位置'), sg.Input(key='file_name', size=(51, None))],
        [sg.FileBrowse('选择文件', target='file_name'), sg.Button('开始检测'),
         sg.CBox('中文分号分句', default=False, key='split_type'), sg.Button('退出')]
    ]
    # 窗口栏
    windows1 = sg.Window('纠错帮V1.1', layout=layout1, font=my_font_style1)
    for i in range(10):
        event1, value1 = windows1.read()
        if event1 in ('退出', None):
            break
        elif event1 == '使用说明':
            sg.popup('1.搜索百度AI开放平台', '2.点击控制台，注册并登录', '3.选择自然语音处理', '4.创建应用',
                     '5.填写appid, api_key, secret_key', '6.选择需要纠错的文件', '7.点击开始检测', title='使用说明',
                     font=my_font_style1)
        elif event1 == '更新记录':
            sg.popup(
                'V1.1更新记录'
                '1.增加了分号分句功能',
                '2.增加了导出word对比功能',
                title='提示', font=my_font_style1)
        elif event1 == '保存配置':
            APP_ID = windows1['app_id'].get()
            API_KEY = windows1['api_key'].get()
            SECRET_KEY = windows1['secret_key'].get()
            file_path = windows1['file_name'].get()
            split_type = windows1['split_type'].get()
            if len(APP_ID) > 3 and len(API_KEY) > 5 and len(SECRET_KEY) > 5:
                dict1 = {
                    'app_id': APP_ID,
                    'api_key': API_KEY,
                    'secret_key': SECRET_KEY,
                    'split_type': split_type
                }
                with open('info.pkl', 'wb') as f:
                    pickle.dump(dict1, f)
                sg.popup('保存完毕', '已经生成一个info.pkl文件到本地', title='提示', auto_close=True,
                         auto_close_duration=3, font=my_font_style1)
            else:
                sg.popup('请检查你的api相关信息是否填写完成', title='错误提示', font=my_font_style1)
        elif event1 == '载入配置':
            if not os.path.exists('info.pkl'):
                sg.popup('没有找到你的配置文件info.pkl', '请检查你的文件是否在当前路径', title='错误提示',
                         font=my_font_style1)
            else:
                with open('info.pkl', 'rb') as f:
                    dict2 = pickle.load(f)
                    windows1['app_id'].update(dict2['app_id'])
                    windows1['api_key'].update(dict2['api_key'])
                    windows1['secret_key'].update(dict2['secret_key'])
                    windows1['split_type'].update(dict2['split_type'])
                    sg.popup('配置文件载入完毕', title='提示', auto_close_duration=3, auto_close=True, font=my_font_style1)
        elif event1 == '开始检测':
            APP_ID = windows1['app_id'].get()
            API_KEY = windows1['api_key'].get()
            SECRET_KEY = windows1['secret_key'].get()
            file_path = windows1['file_name'].get()
            split_type = windows1['split_type'].get()
            if len(APP_ID) > 3 and len(API_KEY) > 5 and len(SECRET_KEY) > 5:
                doc = TextAIAnalyse(file_path, APP_ID, API_KEY, SECRET_KEY)
                text_list2 = doc.text_list2
                if split_type:  # 如果选择的True,也就是支持分号
                    text_list3 = doc.split_text(text_list2)
                else:
                    text_list3 = doc.split_text2(text_list2)
                sg.popup('开始检测，共有{}句'.format(len(text_list3)),
                         '预计用时{}秒'.format(len(text_list3)*2), auto_close_duration=5, auto_close=True)
                layout2 = [
                    [sg.Text('处理进度条', font=my_font_style1),
                     sg.ProgressBar(len(text_list3), orientation='h', key='bar', size=(50, 20))],
                    [sg.Button('取消', font=my_font_style1)]
                ]
                windows2 = sg.Window(title='进度条', layout=layout2, font=my_font_style1)
                bar = windows2['bar']
                for ii in range(len(text_list3)):
                    event2, value2 = windows2.read(timeout=10)
                    if event2 in ('取消', None):
                        break
                    result2 = doc.ai_analyse(text_list3[ii])
                    print(text_list3[ii])
                    if bool(result2):
                        print(result2)
                        doc.save_analyse(result2)
                    time.sleep(0.5 + random.random() / 10)
                    bar.UpdateBar(ii + i)
                windows2.close()
                sg.popup('已经检测完成', '并且生成了一个“分析结果”文件夹到本地', title='提示', font=my_font_style1)
            elif len(file_path) < 5:
                sg.popup('亲！', '你还没有选择检测的文件', title='提示', font=my_font_style1)
            else:
                sg.popup('请输入你的api信息', '详情请查看使用说明', title='提示',font=my_font_style1)
    windows1.close()
